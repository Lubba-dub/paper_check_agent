"""
参考文献引擎 — 解析、生成、交叉验证

核心能力:
1. 解析: 从 BibTeX / LaTeX bibitem / Word 提取参考文献
2. 生成: 按 IEEE/APA/ACM/Nature/Springer 格式输出
3. 交叉验证: 正文引用 ↔ 参考文献表 一致性检查
4. 质量检查: DOI 验证、元数据匹配、时效性分析
"""
from __future__ import annotations
import json
import logging
import re
import urllib.parse
import urllib.request
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime

logger = logging.getLogger(__name__)


# ─── 数据模型 ─────────────────────────────────────────

@dataclass
class Reference:
    """单条参考文献"""
    ref_id: str = ""
    title: str = ""
    authors: List[str] = field(default_factory=list)
    year: Optional[int] = None
    journal: Optional[str] = None
    booktitle: Optional[str] = None
    publisher: Optional[str] = None
    volume: Optional[str] = None
    number: Optional[str] = None
    pages: Optional[str] = None
    doi: Optional[str] = None
    url: Optional[str] = None
    bibtex_type: str = "article"  # article/inproceedings/book/phdthesis/misc
    source: str = ""              # 从哪提取的: bibtex/bibitem/docx
    raw_text: str = ""            # 原始文本

    # 验证状态
    verified: Optional[bool] = None
    verification_source: str = ""

    def format_ieee(self) -> str:
        """IEEE 格式输出"""
        authors_str = ", ".join(self._format_author_ieee(a) for a in self.authors[:6])
        if len(self.authors) > 6:
            authors_str += " et al."
        title_str = self.title
        if self.bibtex_type == "article" and self.journal:
            return f'{authors_str}, "{title_str}," {self.journal}, vol. {self.volume or ""}, no. {self.number or ""}, pp. {self.pages or ""}, {self.year}.'
        elif self.bibtex_type == "inproceedings" and self.booktitle:
            return f'{authors_str}, "{title_str}," in {self.booktitle}, pp. {self.pages or ""}, {self.year}.'
        elif self.bibtex_type == "book":
            return f'{authors_str}, {title_str}, {self.publisher or ""}, {self.year}.'
        return f'{authors_str}, "{title_str}," {self.year}.'

    def format_apa(self) -> str:
        """APA 格式输出"""
        authors_str = ", ".join(self._format_author_apa(a) for a in self.authors[:7])
        if len(self.authors) > 7:
            authors_str += " ... " + self._format_author_apa(self.authors[-1])
        year_str = f"({self.year})" if self.year else "(n.d.)"
        title_str = self.title
        if self.bibtex_type == "article" and self.journal:
            return f'{authors_str} {year_str}. {title_str}. {self.journal}, {self.volume or ""}({self.number or ""}), {self.pages or ""}. https://doi.org/{self.doi}' if self.doi else f'{authors_str} {year_str}. {title_str}. {self.journal}, {self.volume or ""}({self.number or ""}), {self.pages or ""}.'
        return f'{authors_str} {year_str}. {title_str}.'

    def format_acm(self) -> str:
        """ACM 格式输出"""
        authors_str = ", ".join(self._format_author_acm(a) for a in self.authors[:6])
        if len(self.authors) > 6:
            authors_str += " et al."
        if self.bibtex_type == "article" and self.journal:
            return f'{authors_str}. {self.year}. {self.title}. {self.journal} {self.volume or ""}, {self.number or ""} ({self.year}), {self.pages or ""}.'
        return f'{authors_str}. {self.year}. {self.title}.'

    def _format_author_ieee(self, name: str) -> str:
        parts = name.strip().split(",")
        if len(parts) == 2:
            return f"{parts[1].strip()[0]}. {parts[0].strip()}"  # "Smith, John" → "J. Smith"
        parts2 = name.strip().split()
        if len(parts2) >= 2:
            return f"{parts2[-1][0]}. {' '.join(parts2[:-1])}"
        return name

    def _format_author_apa(self, name: str) -> str:
        parts = name.strip().split(",")
        if len(parts) == 2:
            return f"{parts[0].strip()}, {parts[1].strip()[0]}."
        return name

    def _format_author_acm(self, name: str) -> str:
        parts = name.strip().split(",")
        if len(parts) == 2:
            return f"{parts[1].strip()[0]}. {parts[0].strip()}"
        return name


@dataclass
class Citation:
    """正文中的引用"""
    text: str          # 如 "[1]" 或 "(Smith, 2020)"
    ref_ids: List[str] = field(default_factory=list)
    page: Optional[int] = None
    context_before: str = ""   # 引用前的文本
    context_after: str = ""    # 引用后的文本


@dataclass
class ReferenceCheckResult:
    """文献检查结果"""
    total_refs: int = 0
    total_citations: int = 0
    matched: int = 0              # 引用 → 参考文献 匹配数
    unmatched_citations: List[str] = field(default_factory=list)  # 正文引用但不存在于参考文献
    unused_refs: List[str] = field(default_factory=list)          # 参考文献存在但正文未引用
    format_issues: List[Dict] = field(default_factory=list)
    doi_missing: List[str] = field(default_factory=list)
    score: float = 1.0


# ═══════════════════════════════════════════════════════
# 解析器
# ═══════════════════════════════════════════════════════

class ReferenceParser:
    """文献解析器 — 从各种来源提取参考文献"""

    @staticmethod
    def from_bibtex(text: str) -> List[Reference]:
        """从 BibTeX 文本解析"""
        refs = []
        # 匹配每个 @type{key, ...}
        entries = re.finditer(
            r'@(\w+)\{(\w+),\s*([^@]+)\}',
            text, re.DOTALL
        )
        for m in entries:
            ref = Reference(ref_id=m.group(2), bibtex_type=m.group(1).lower())
            body = m.group(3)
            # 提取字段
            fields = re.findall(r'(\w+)\s*=\s*\{([^}]+)\}', body)
            fields += re.findall(r'(\w+)\s*=\s*"([^"]+)"', body)
            for key, value in fields:
                key = key.lower()
                if key == "title":
                    ref.title = value.strip("{}")
                elif key == "author":
                    ref.authors = [a.strip() for a in value.split(" and ")]
                elif key == "year":
                    ref.year = int(re.search(r'\d{4}', value).group()) if re.search(r'\d{4}', value) else None
                elif key == "journal":
                    ref.journal = value.strip("{}")
                elif key == "booktitle":
                    ref.booktitle = value.strip("{}")
                elif key == "publisher":
                    ref.publisher = value.strip("{}")
                elif key == "volume":
                    ref.volume = value.strip("{}")
                elif key == "number":
                    ref.number = value.strip("{}")
                elif key == "pages":
                    ref.pages = value.strip("{}")
                elif key == "doi":
                    ref.doi = value.strip("{}")
                elif key == "url":
                    ref.url = value.strip("{}")
            ref.source = "bibtex"
            refs.append(ref)
        return refs

    @staticmethod
    def from_bibitem(text: str) -> List[Reference]:
        """从 LaTeX \\bibitem 解析"""
        refs = []
        items = re.finditer(
            r'\\bibitem(?:\[([^\]]*)\])?\{(.+?)\}\s*(.+?)(?=\\bibitem|\Z)',
            text, re.DOTALL
        )
        for m in items:
            ref = Reference(ref_id=m.group(2), raw_text=m.group(3).strip())
            ref.source = "bibitem"
            # 解析原始文本
            ref.title = ref.raw_text[:100]
            # 尝试提取年份
            year_match = re.search(r'(19|20)\d{2}', ref.raw_text)
            if year_match:
                ref.year = int(year_match.group())
            refs.append(ref)
        return refs

    @staticmethod
    def from_docx(doc_path: str) -> List[Reference]:
        """从 Word 文档参考文献段落提取（支持多行条目合并与类型识别）"""
        refs = []
        try:
            from docx import Document
            doc = Document(doc_path)
            entries = ReferenceParser._extract_docx_reference_entries(doc)
            for index, entry in enumerate(entries, start=1):
                ref = ReferenceParser._parse_reference_entry(entry, index=index)
                if ref:
                    refs.append(ref)
            return refs
        except Exception as e:
            logger.error(f"Word 文献提取失败: {e}")
            return []

    @staticmethod
    def extract_body_text(doc_path: str) -> str:
        """提取正文文本，尽量排除参考文献区，避免把文献编号当成正文引文。"""
        try:
            from docx import Document
            doc = Document(doc_path)
        except Exception:
            return ""

        body_lines: List[str] = []
        in_refs = False
        stop_norms = {"致谢", "附录", "appendix", "acknowledgements", "acknowledgments"}

        for block in ReferenceParser._iter_docx_text_blocks(doc):
            text = block.strip()
            if not text:
                continue
            normalized = ReferenceParser._normalize_inline_text(text).rstrip(":：")
            if normalized in {"参考文献", "references", "bibliography"}:
                in_refs = True
                continue
            if in_refs:
                if normalized in stop_norms:
                    in_refs = False
                continue
            body_lines.append(text)
        return "\n".join(body_lines)

    @staticmethod
    def _iter_docx_text_blocks(doc) -> List[str]:
        para_map = {id(para._p): para for para in doc.paragraphs}
        table_map = {id(table._tbl): table for table in doc.tables}
        blocks: List[str] = []

        for child in doc.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p" and id(child) in para_map:
                text = para_map[id(child)].text.strip()
                if text:
                    blocks.append(text)
            elif tag == "tbl" and id(child) in table_map:
                table = table_map[id(child)]
                for row in table.rows:
                    row_cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()) for cell in row.cells]
                    row_text = " ".join(item.strip() for item in row_cells if item and item.strip()).strip()
                    if row_text:
                        blocks.append(row_text)
        return blocks

    @staticmethod
    def _extract_docx_reference_entries(doc) -> List[str]:
        lines = ReferenceParser._iter_docx_text_blocks(doc)
        in_refs = False
        collected: List[str] = []

        for line in lines:
            normalized = ReferenceParser._normalize_inline_text(line).rstrip(":：")
            if normalized in {"参考文献", "references", "bibliography"}:
                in_refs = True
                continue
            if not in_refs:
                continue
            if normalized in {"致谢", "附录", "appendix", "acknowledgements", "acknowledgments"}:
                break
            if len(line.strip()) < 2:
                continue
            collected.append(line.strip())

        entries: List[str] = []
        current = ""
        entry_start = re.compile(r"^\s*(\[(\d+)\]|(\d+)[.、])\s*")
        for line in collected:
            if entry_start.match(line):
                if current:
                    entries.append(current.strip())
                current = line.strip()
            elif current:
                current = f"{current} {line.strip()}".strip()
            else:
                current = line.strip()
        if current:
            entries.append(current.strip())
        return entries

    @staticmethod
    def _parse_reference_entry(text: str, *, index: int) -> Optional[Reference]:
        if not text:
            return None

        ref = Reference(raw_text=text.strip(), source="docx")
        num_match = re.match(r"^\s*(?:\[(\d+)\]|(\d+)[.、])\s*", text)
        numeric_id = next((group for group in num_match.groups() if group), None) if num_match else None
        ref.ref_id = numeric_id or str(index)
        body = re.sub(r"^\s*(?:\[\d+\]|\d+[.、])\s*", "", text).strip()

        doi_match = re.search(r"(10\.\d{4,}/[-._;()/:A-Za-z0-9]+)", body, re.IGNORECASE)
        if doi_match:
            ref.doi = doi_match.group(1)
        url_match = re.search(r"(https?://\S+)", body, re.IGNORECASE)
        if url_match:
            ref.url = url_match.group(1)
        year_match = re.search(r"(19|20)\d{2}", body)
        if year_match:
            ref.year = int(year_match.group())

        ref.bibtex_type = ReferenceValidator.infer_reference_type(
            raw_text=body,
            title="",
            journal="",
            booktitle="",
            publisher="",
        )

        sentence_parts = [part.strip(" .。；;") for part in re.split(r"[.。]\s*", body) if part.strip(" .。；;")]
        if sentence_parts:
            author_text = sentence_parts[0]
            author_parts = re.split(r"[，,、;；]\s*", author_text)
            ref.authors = [item.strip() for item in author_parts if len(item.strip()) > 1][:10]
        if len(sentence_parts) >= 2:
            ref.title = sentence_parts[1][:240]
        else:
            title_match = re.search(r"(?:(?:19|20)\d{2}[,，]?\s*)?(.+?)(?:\[?[JMDBSENCP/OL]+\]?|$)", body)
            ref.title = (title_match.group(1).strip() if title_match else body[:240]).strip(" ,.;。；")

        if ref.bibtex_type == "article":
            journal_match = re.search(r"\[\s*J\s*\][,，.\s]*(.+?)(?:[,，]\s*(?:19|20)\d{2}|$)", body, re.IGNORECASE)
            if journal_match:
                ref.journal = journal_match.group(1).strip()
        elif ref.bibtex_type == "book":
            publisher_match = re.search(r"(出版社|Press)", body, re.IGNORECASE)
            if publisher_match:
                ref.publisher = body[max(0, publisher_match.start() - 24): publisher_match.end() + 24].strip(" ,.;。；")

        return ref

    @staticmethod
    def _normalize_inline_text(value: str) -> str:
        return re.sub(r"\s+", "", str(value or "")).lower()

    @staticmethod
    def extract_citations(text: str, style: str = "numeric") -> List[Citation]:
        """从正文中提取引用"""
        citations = []
        if style == "numeric":
            # [1], [1,2], [1-3]
            matches = re.finditer(r'\[([\d,\-\s]+)\]', text)
            for m in matches:
                ref_ids = re.findall(r'\d+', m.group(1))
                citations.append(Citation(text=m.group(), ref_ids=ref_ids))
        else:
            # (Author, Year) 或 Author (Year)
            matches = re.finditer(r'\(([^)]+\d{4}[^)]*)\)', text)
            for m in matches:
                citations.append(Citation(text=m.group(), context_before=text[max(0,m.start()-30):m.start()]))
        return citations


# ═══════════════════════════════════════════════════════
# 生成器
# ═══════════════════════════════════════════════════════

class ReferenceGenerator:
    """参考文献生成器 — 按格式输出"""

    FORMATS = {
        "ieee": {"name": "IEEE", "citation": "[{}]", "ref_prefix": "[{}] "},
        "apa": {"name": "APA", "citation": "({}, {})", "ref_prefix": "{} "},
        "acm": {"name": "ACM", "citation": "[{} {}]", "ref_prefix": "[{}] "},
        "springer": {"name": "Springer LNCS", "citation": "[{}]", "ref_prefix": "[{}] "},
        "nature": {"name": "Nature", "citation": "^{}", "ref_prefix": "{}. "},
    }

    @staticmethod
    def format_reference(ref: Reference, style: str = "ieee", index: int = 1) -> str:
        """格式化单条参考文献"""
        fmt = ReferenceGenerator.FORMATS.get(style, ReferenceGenerator.FORMATS["ieee"])
        prefix = fmt["ref_prefix"].format(index)

        if style == "ieee":
            return prefix + ref.format_ieee()
        elif style == "apa":
            return prefix + ref.format_apa()
        elif style == "acm":
            return prefix + ref.format_acm()
        return prefix + ref.format_ieee()

    @staticmethod
    def format_bibliography(refs: List[Reference], style: str = "ieee") -> str:
        """生成完整的参考文献列表"""
        lines = ["## 参考文献\n"]
        for i, ref in enumerate(refs, 1):
            lines.append(ReferenceGenerator.format_reference(ref, style, i))
            lines.append("")
        return "\n".join(lines)

    @staticmethod
    def format_inline_citation(ref_ids: List[str], style: str = "ieee") -> str:
        """生成正文引用标记"""
        if style == "ieee":
            return f"[{','.join(ref_ids)}]"
        elif style == "acm":
            return f"[{', '.join(ref_ids)}]"
        elif style == "apa":
            return f"({', '.join(ref_ids)})"
        return f"[{','.join(ref_ids)}]"


# ═══════════════════════════════════════════════════════
# 交叉验证
# ═══════════════════════════════════════════════════════

class ReferenceValidator:
    """文献验证器 — 交叉检查引用一致性"""

    @staticmethod
    def check_consistency(
        refs: List[Reference],
        citations: List[Citation],
        text: str,
    ) -> ReferenceCheckResult:
        """检查正文引用与参考文献表的一致性"""
        result = ReferenceCheckResult(
            total_refs=len(refs),
            total_citations=len(citations),
        )

        # 收集所有引文ID
        cited_ids = set()
        for cit in citations:
            for rid in cit.ref_ids:
                cited_ids.add(rid)

        # 收集所有参考文献ID
        ref_ids = {str(i+1) for i in range(len(refs))}
        ref_ids |= {r.ref_id for r in refs if r.ref_id}

        # 未匹配的引用
        for rid in cited_ids:
            if rid not in ref_ids:
                result.unmatched_citations.append(rid)

        # 未使用的参考文献
        ref_id_by_index = {str(i+1): r for i, r in enumerate(refs)}
        for rid, ref in ref_id_by_index.items():
            if rid not in cited_ids:
                result.unused_refs.append(rid)

        result.matched = len(cited_ids & ref_ids)

        # 检查 DOI 缺失
        for ref in refs:
            if not ref.doi and ReferenceValidator.requires_doi(ref):
                result.doi_missing.append(ref.ref_id or f"ref_{refs.index(ref)+1}")

        result.score = ReferenceValidator._compute_score(result)
        return result

    @staticmethod
    def _compute_score(result: ReferenceCheckResult) -> float:
        """计算文献一致性评分"""
        score = 1.0
        if result.total_citations > 0:
            unmatched_ratio = len(result.unmatched_citations) / result.total_citations
            score -= unmatched_ratio * 0.5
        if result.total_refs > 0:
            unused_ratio = len(result.unused_refs) / result.total_refs
            score -= unused_ratio * 0.3
            doi_missing_ratio = len(result.doi_missing) / result.total_refs
            score -= doi_missing_ratio * 0.2
        return max(0.0, round(score, 2))

    @staticmethod
    def infer_reference_type(
        raw_text: str,
        title: str = "",
        journal: str = "",
        booktitle: str = "",
        publisher: str = "",
    ) -> str:
        text = " ".join(str(item or "") for item in [raw_text, title, journal, booktitle, publisher])
        upper = text.upper()
        if re.search(r"(GB/T|ISO|IEC|标准|国家标准|\[S\])", upper, re.IGNORECASE):
            return "standard"
        if re.search(r"(学位论文|硕士学位论文|博士学位论文|\[D\])", text, re.IGNORECASE):
            return "phdthesis"
        if re.search(r"(\[M\]|出版社|PRESS\b|出版)", text, re.IGNORECASE):
            return "book"
        if re.search(r"(\[EB/OL\]|\[EB\]|\[OL\]|https?://|www\.)", text, re.IGNORECASE):
            return "webpage"
        if journal or re.search(r"(\[J\]|journal|学报|期刊)", text, re.IGNORECASE):
            return "article"
        if booktitle or re.search(r"(\[C\]|conference|proceedings|会议)", text, re.IGNORECASE):
            return "inproceedings"
        return "misc"

    @staticmethod
    def requires_doi(ref: Reference) -> bool:
        ref_type = ref.bibtex_type or ReferenceValidator.infer_reference_type(
            raw_text=ref.raw_text,
            title=ref.title,
            journal=ref.journal or "",
            booktitle=ref.booktitle or "",
            publisher=ref.publisher or "",
        )
        return ref_type in {"article", "inproceedings"}

    @staticmethod
    def requires_doi_metadata(ref: Dict[str, Any]) -> bool:
        ref_type = str(ref.get("bibtex_type") or ref.get("reference_type") or "").strip().lower()
        if not ref_type:
            ref_type = ReferenceValidator.infer_reference_type(
                raw_text=str(ref.get("raw_text") or ""),
                title=str(ref.get("title") or ""),
                journal=str(ref.get("journal") or ""),
                booktitle=str(ref.get("booktitle") or ""),
                publisher=str(ref.get("publisher") or ""),
            )
        return ref_type in {"article", "inproceedings"}

    @staticmethod
    def verify_doi(doi: str) -> Dict[str, Any]:
        """验证 DOI（使用 CrossRef API）"""
        try:
            url = f"https://api.crossref.org/works/{doi}"
            req = urllib.request.Request(url, headers={"User-Agent": "ArticleCheck/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode())
                msg = data.get("message", {})
                return {
                    "valid": True,
                    "title": msg.get("title", [""])[0] if msg.get("title") else "",
                    "authors": [a.get("given", "") + " " + a.get("family", "") for a in msg.get("author", [])],
                    "year": msg.get("published-print", {}).get("date-parts", [[None]])[0][0]
                            or msg.get("created", {}).get("date-parts", [[None]])[0][0],
                    "journal": msg.get("container-title", [""])[0] if msg.get("container-title") else "",
                }
        except Exception as e:
            return {"valid": False, "error": str(e)}

    @staticmethod
    def check_reference_exists(
        title: str,
        authors: Optional[List[str]] = None,
        year: Optional[int] = None,
    ) -> Dict[str, Any]:
        """通过 Crossref / OpenAlex 检查参考文献题名是否真实存在。"""
        normalized_title = ReferenceValidator._normalize_title(title)
        if not normalized_title:
            return {"exists": False, "message": "题名为空，无法检索"}

        crossref_result = ReferenceValidator._search_crossref(title, authors=authors, year=year)
        if crossref_result.get("exists"):
            return crossref_result

        openalex_result = ReferenceValidator._search_openalex(title, authors=authors, year=year)
        if openalex_result.get("exists"):
            return openalex_result

        return {
            "exists": False,
            "query_title": title,
            "message": openalex_result.get("message") or crossref_result.get("message") or "未在 Crossref/OpenAlex 中检索到高置信候选条目",
            "sources": [crossref_result, openalex_result],
        }

    @staticmethod
    def _search_crossref(title: str, authors: Optional[List[str]] = None, year: Optional[int] = None) -> Dict[str, Any]:
        try:
            params = urllib.parse.urlencode({"query.bibliographic": title, "rows": 3})
            url = f"https://api.crossref.org/works?{params}"
            req = urllib.request.Request(url, headers={"User-Agent": "ArticleCheck/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode())
            items = (data.get("message") or {}).get("items") or []
            return ReferenceValidator._pick_best_candidate("crossref", title, items, authors=authors, year=year)
        except Exception as exc:
            return {"exists": False, "source": "crossref", "message": str(exc)}

    @staticmethod
    def _search_openalex(title: str, authors: Optional[List[str]] = None, year: Optional[int] = None) -> Dict[str, Any]:
        try:
            params = urllib.parse.urlencode({"search": title, "per-page": 3})
            url = f"https://api.openalex.org/works?{params}"
            req = urllib.request.Request(url, headers={"User-Agent": "ArticleCheck/1.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode())
            items = data.get("results") or []
            return ReferenceValidator._pick_best_candidate("openalex", title, items, authors=authors, year=year)
        except Exception as exc:
            return {"exists": False, "source": "openalex", "message": str(exc)}

    @staticmethod
    def _pick_best_candidate(
        source: str,
        title: str,
        items: List[Dict[str, Any]],
        *,
        authors: Optional[List[str]] = None,
        year: Optional[int] = None,
    ) -> Dict[str, Any]:
        normalized_query = ReferenceValidator._normalize_title(title)
        best: Optional[Dict[str, Any]] = None
        best_score = 0.0

        for item in items:
            candidate_title = ""
            candidate_authors: List[str] = []
            candidate_year = None
            candidate_doi = None

            if source == "crossref":
                candidate_title = ((item.get("title") or [""])[:1] or [""])[0]
                candidate_authors = [
                    " ".join(part for part in [author.get("given", ""), author.get("family", "")] if part).strip()
                    for author in (item.get("author") or [])
                ]
                candidate_year = (item.get("published-print") or item.get("published-online") or item.get("created") or {}).get("date-parts", [[None]])[0][0]
                candidate_doi = item.get("DOI")
            else:
                candidate_title = item.get("display_name") or ""
                candidate_authors = [
                    (author.get("author") or {}).get("display_name", "")
                    for author in (item.get("authorships") or [])
                ]
                candidate_year = item.get("publication_year")
                candidate_doi = item.get("doi")

            title_score = ReferenceValidator._title_similarity(normalized_query, ReferenceValidator._normalize_title(candidate_title))
            author_score = ReferenceValidator._author_overlap(authors or [], candidate_authors)
            year_score = 1.0 if not year or not candidate_year or int(year) == int(candidate_year) else 0.0
            score = title_score * 0.75 + author_score * 0.15 + year_score * 0.10

            if score > best_score:
                best_score = score
                best = {
                    "exists": score >= 0.72,
                    "source": source,
                    "query_title": title,
                    "matched_title": candidate_title,
                    "matched_authors": candidate_authors[:6],
                    "matched_year": candidate_year,
                    "matched_doi": candidate_doi,
                    "title_similarity": round(title_score, 3),
                    "score": round(score, 3),
                }

        return best or {"exists": False, "source": source, "query_title": title, "message": "未找到候选条目"}

    @staticmethod
    def _normalize_title(title: str) -> str:
        normalized = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", (title or "").lower())
        return normalized.strip()

    @staticmethod
    def _title_similarity(a: str, b: str) -> float:
        if not a or not b:
            return 0.0
        if a == b:
            return 1.0
        shared = len(set(a) & set(b))
        return shared / max(len(set(a) | set(b)), 1)

    @staticmethod
    def _author_overlap(left: List[str], right: List[str]) -> float:
        left_norm = {ReferenceValidator._normalize_title(item) for item in left if item}
        right_norm = {ReferenceValidator._normalize_title(item) for item in right if item}
        if not left_norm or not right_norm:
            return 0.0
        return len(left_norm & right_norm) / max(len(left_norm), 1)


# ═══════════════════════════════════════════════════════
# 高层接口
# ═══════════════════════════════════════════════════════

class ReferenceEngine:
    """文献引擎统一入口"""

    def __init__(self):
        self.parser = ReferenceParser()
        self.generator = ReferenceGenerator()
        self.validator = ReferenceValidator()
        self._cached_refs: List[Reference] = []
        logger.info("ReferenceEngine 初始化完成")

    def extract_from_paper(self, paper_path: str) -> List[Reference]:
        """从论文文件中智能提取参考文献"""
        path = Path(paper_path)
        suffix = path.suffix.lower()

        if suffix in (".tex", ".ltx"):
            text = path.read_text(encoding="utf-8", errors="replace")
            # 优先 BibTeX
            if ".bib" in text:
                bib_files = re.findall(r'\\addbibresource\{(.+?)\}', text)
                bib_files += re.findall(r'\\bibliography\{(.+?)\}', text)
                refs = []
                for bf in bib_files:
                    bib_path = path.parent / (bf if bf.endswith(".bib") else bf + ".bib")
                    if bib_path.exists():
                        refs.extend(self.parser.from_bibtex(bib_path.read_text(encoding="utf-8", errors="replace")))
                if refs:
                    self._cached_refs = refs
                    return refs
            # 或 bibitem
            refs = self.parser.from_bibitem(text)
            self._cached_refs = refs
            return refs

        elif suffix == ".docx":
            refs = self.parser.from_docx(paper_path)
            self._cached_refs = refs
            return refs

        elif suffix == ".bib":
            refs = self.parser.from_bibtex(path.read_text(encoding="utf-8", errors="replace"))
            self._cached_refs = refs
            return refs

        return []

    def generate_bibliography(
        self,
        refs: Optional[List[Reference]] = None,
        style: str = "ieee",
    ) -> str:
        """生成格式化参考文献列表"""
        if refs is None:
            refs = self._cached_refs
        return self.generator.format_bibliography(refs, style)

    def validate(
        self,
        paper_path: str,
        refs: Optional[List[Reference]] = None,
    ) -> ReferenceCheckResult:
        """论文文献一致性验证"""
        if refs is None:
            refs = self.extract_from_paper(paper_path)

        # 提取正文
        path = Path(paper_path)
        if path.suffix.lower() in (".tex", ".ltx"):
            text = path.read_text(encoding="utf-8", errors="replace")
            citations = self.parser.extract_citations(text, "numeric")
        elif path.suffix.lower() == ".docx":
            text = self.parser.extract_body_text(paper_path)
            citations = self.parser.extract_citations(text, "numeric")
        else:
            return ReferenceCheckResult()

        return self.validator.check_consistency(refs, citations, text)

    def check_ref_quality(self, ref: Reference) -> Dict[str, Any]:
        """单条文献质量检查"""
        result = {
            "ref_id": ref.ref_id,
            "title": ref.title[:60] if ref.title else "❌ 缺失",
            "authors": len(ref.authors),
            "year": ref.year or "❌ 缺失",
            "has_doi": bool(ref.doi),
        }
        if ref.doi:
            doi_check = self.validator.verify_doi(ref.doi)
            result["doi_verified"] = doi_check.get("valid", False)
            result["doi_details"] = doi_check
        if ref.title:
            exists_check = self.validator.check_reference_exists(ref.title, authors=ref.authors, year=ref.year)
            result["exists"] = exists_check.get("exists", False)
            result["existence_details"] = exists_check
        return result

    def get_report(self) -> Dict[str, Any]:
        """获取文献引擎状态报告"""
        return {
            "cached_refs": len(self._cached_refs),
            "available_formats": list(ReferenceGenerator.FORMATS.keys()),
            "engine_ready": True,
        }
