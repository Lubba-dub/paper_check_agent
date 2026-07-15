from __future__ import annotations

import logging
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from article_check.config.settings import config
from article_check.parsers import GrobidClient
from article_check.references.engine import ReferenceEngine, ReferenceParser, ReferenceValidator
from article_check.utils.file_utils import detect_file_type, extract_text_from_docx, extract_text_from_pdf, read_paper_content

logger = logging.getLogger(__name__)


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _detect_language(text: str) -> str:
    if not text:
        return ""
    zh_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    en_count = len(re.findall(r"[A-Za-z]", text))
    return "zh" if zh_count >= en_count else "en"


def _detect_heading_level(line: str) -> Optional[int]:
    candidate = line.strip()
    if not candidate or len(candidate) > 100:
        return None

    normalized = _normalize_text(candidate).rstrip(":：")
    if normalized in {
        "摘要", "abstract", "关键词", "目录", "引言", "前言", "绪论", "正文",
        "结论", "结语", "参考文献", "致谢", "附录", "references", "bibliography",
    }:
        return 1
    if re.match(r"^第[一二三四五六七八九十\d]+章", candidate):
        return 1
    if re.match(r"^[一二三四五六七八九十]+[、.．].+", candidate):
        return 1
    if re.match(r"^\d+\.\d+(\.\d+){0,3}\s*\S+", candidate):
        return min(candidate.count(".") + 1, 4)
    if re.match(r"^\d+[.、]\s*\S+", candidate):
        return 2
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]\s*\S+", candidate):
        return 3
    return None


def _looks_like_caption(text: str) -> bool:
    return bool(re.match(r"^(图|表)\s*\d+|^(Figure|Table)\s+\d+", str(text or "").strip(), re.IGNORECASE))


def _detect_review_track(text: str, *, template_name: str = "", paper_path: Optional[Path] = None) -> Tuple[str, List[str]]:
    joined = " ".join([template_name or "", str(paper_path or ""), text[:5000]])
    score = {"undergraduate": 0, "graduate": 0}
    reasons: List[str] = []

    undergraduate_markers = ["本科", "毕业论文", "本科生"]
    graduate_markers = ["研究生", "硕士", "博士", "学位论文", "题名页", "版权页"]

    for marker in undergraduate_markers:
        if marker in joined:
            score["undergraduate"] += 2
            reasons.append(f"命中本科标记: {marker}")
    for marker in graduate_markers:
        if marker in joined:
            score["graduate"] += 2
            reasons.append(f"命中研究生标记: {marker}")

    if score["graduate"] > score["undergraduate"]:
        return "graduate", reasons
    if score["undergraduate"] > score["graduate"]:
        return "undergraduate", reasons
    return "graduate", reasons or ["未命中明确标记，默认按研究生规则高标准审查"]


def _infer_title(path: Path, text: str) -> str:
    for line in text.splitlines():
        candidate = line.strip()
        if candidate and len(candidate) <= 80 and not _detect_heading_level(candidate):
            return candidate
    return path.stem


def _serialize_reference(ref: Any, index: int) -> Dict[str, Any]:
    return {
        "reference_id": ref.ref_id or str(index),
        "anchor_id": f"reference-{index}",
        "title": ref.title,
        "authors": list(ref.authors or []),
        "year": ref.year,
        "journal": ref.journal,
        "booktitle": ref.booktitle,
        "publisher": ref.publisher,
        "volume": ref.volume,
        "number": ref.number,
        "pages": ref.pages,
        "doi": ref.doi,
        "arxiv_id": getattr(ref, "arxiv_id", None),
        "pmid": getattr(ref, "pmid", None),
        "url": ref.url,
        "bibtex_type": getattr(ref, "bibtex_type", None),
        "source": ref.source,
        "raw_text": ref.raw_text,
    }


def _serialize_reference_stats(result: Any) -> Dict[str, Any]:
    if not result:
        return {
            "total_refs": 0,
            "total_citations": 0,
            "matched": 0,
            "unmatched_citations": [],
            "unused_refs": [],
            "doi_missing": [],
            "score": 0.0,
        }
    return {
        "total_refs": getattr(result, "total_refs", 0),
        "total_citations": getattr(result, "total_citations", 0),
        "matched": getattr(result, "matched", 0),
        "unmatched_citations": getattr(result, "unmatched_citations", []),
        "unused_refs": getattr(result, "unused_refs", []),
        "doi_missing": getattr(result, "doi_missing", []),
        "parsed_entry_count": getattr(result, "total_refs", 0),
        "score": getattr(result, "score", 0.0),
    }


def _extract_citation_contexts(text: str, radius: int = 120) -> List[Dict[str, Any]]:
    citations: List[Dict[str, Any]] = []
    pattern = re.compile(r"\[([\d,\-\s]+)\]")
    for index, match in enumerate(pattern.finditer(text), start=1):
        ref_ids = re.findall(r"\d+", match.group(1))
        citations.append(
            {
                "citation_id": f"citation-{index}",
                "text": match.group(0),
                "ref_ids": ref_ids,
                "context_before": text[max(0, match.start() - radius): match.start()].strip(),
                "context_after": text[match.end(): min(len(text), match.end() + radius)].strip(),
                "char_start": match.start(),
                "char_end": match.end(),
            }
        )
    return citations


def _compute_reference_stats_from_records(references: List[Dict[str, Any]], citations: List[Dict[str, Any]]) -> Dict[str, Any]:
    ref_ids = {
        str(item.get("reference_id"))
        for item in references
        if item.get("reference_id") not in (None, "")
    }
    cited_ids = {
        str(ref_id)
        for citation in citations
        for ref_id in (citation.get("ref_ids") or [])
        if str(ref_id).strip()
    }
    unmatched = sorted(cited_ids - ref_ids)
    unused = sorted(ref_ids - cited_ids)
    doi_missing = [
        item.get("reference_id") or f"ref_{index}"
        for index, item in enumerate(references, start=1)
        if not item.get("doi") and ReferenceValidator.requires_doi_metadata(item)
    ]
    total_refs = len(references)
    total_citations = len(citations)
    matched = len(cited_ids & ref_ids)
    score = 1.0
    if total_citations:
        score -= (len(unmatched) / total_citations) * 0.5
    if total_refs:
        score -= (len(unused) / total_refs) * 0.3
        score -= (len(doi_missing) / total_refs) * 0.2
    return {
        "total_refs": total_refs,
        "total_citations": total_citations,
        "matched": matched,
        "unmatched_citations": unmatched,
        "unused_refs": unused,
        "doi_missing": doi_missing,
        "parsed_entry_count": total_refs,
        "score": max(0.0, round(score, 3)),
    }


def _extract_section_text_from_blocks(
    blocks: List[Dict[str, Any]],
    start_aliases: List[str],
    stop_aliases: List[str],
) -> str:
    normalized_starts = {_normalize_text(item).rstrip(":：") for item in start_aliases}
    normalized_stops = {_normalize_text(item).rstrip(":：") for item in stop_aliases}
    in_section = False
    collected: List[str] = []

    for block in blocks:
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        normalized = _normalize_text(text).rstrip(":：")
        if normalized in normalized_starts:
            in_section = True
            continue
        if in_section and normalized in normalized_stops:
            break
        if in_section:
            collected.append(text)
    return "\n".join(collected).strip()


def _merge_grobid_pdf_parse(base: Dict[str, Any], grobid: Dict[str, Any]) -> Dict[str, Any]:
    if not grobid.get("enabled"):
        return base
    merged = dict(base)
    if grobid.get("blocks"):
        merged["blocks"] = grobid.get("blocks")
    if grobid.get("anchors"):
        merged["anchors"] = grobid.get("anchors")
    if grobid.get("figures"):
        merged["figures"] = grobid.get("figures")
    if grobid.get("tables"):
        merged["tables"] = grobid.get("tables")
    if grobid.get("captions"):
        merged["captions"] = grobid.get("captions")
    merged["grobid"] = grobid
    if grobid.get("title"):
        merged["title"] = grobid.get("title")
    if grobid.get("abstract"):
        merged["abstract"] = grobid.get("abstract")
    if not grobid.get("text") and base.get("text"):
        merged["text"] = base["text"]
    return merged


def _build_structure_from_blocks(blocks: List[Dict[str, Any]], text: str) -> List[Dict[str, Any]]:
    structure: List[Dict[str, Any]] = []
    current_node: Optional[Dict[str, Any]] = None

    for block in blocks:
        if block.get("type") == "heading":
            current_node = {
                "node_id": f"section-{len(structure) + 1}",
                "heading": block.get("text", ""),
                "level": block.get("heading_level", 1),
                "page_start": block.get("page"),
                "page_end": block.get("page"),
                "block_ids": [block.get("block_id")],
                "paragraph_start": block.get("paragraph_index"),
                "paragraph_end": block.get("paragraph_index"),
                "text_excerpt": "",
            }
            structure.append(current_node)
            continue

        if current_node is None:
            current_node = {
                "node_id": "section-0",
                "heading": "全文",
                "level": 0,
                "page_start": block.get("page"),
                "page_end": block.get("page"),
                "block_ids": [block.get("block_id")],
                "paragraph_start": block.get("paragraph_index"),
                "paragraph_end": block.get("paragraph_index"),
                "text_excerpt": "",
            }
            structure.append(current_node)
        else:
            current_node["block_ids"].append(block.get("block_id"))
            current_node["page_end"] = block.get("page") or current_node.get("page_end")
            current_node["paragraph_end"] = block.get("paragraph_index") or current_node.get("paragraph_end")

        excerpt = (current_node.get("text_excerpt") or "")
        if len(excerpt) < 800 and block.get("text"):
            current_node["text_excerpt"] = (excerpt + "\n" + block["text"]).strip()[:800]

    if not structure and text.strip():
        structure.append(
            {
                "node_id": "section-0",
                "heading": "全文",
                "level": 0,
                "page_start": 1,
                "page_end": 1,
                "block_ids": [],
                "paragraph_start": 1,
                "paragraph_end": None,
                "text_excerpt": text[:800],
            }
        )
    return structure


def _read_text(path: Path, file_type: str) -> str:
    if file_type == "docx":
        return extract_text_from_docx(path)
    if file_type == "pdf":
        return extract_text_from_pdf(path)
    return read_paper_content(path)


def _parse_docx_layout(path: Path) -> Dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    blocks: List[Dict[str, Any]] = []
    anchors: List[Dict[str, Any]] = []
    captions: List[Dict[str, Any]] = []
    figures: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []
    font_sizes: List[float] = []
    font_names: List[str] = []
    line_spacings: List[float] = []
    para_map = {id(para._p): para for para in doc.paragraphs}
    table_map = {id(table._tbl): table for table in doc.tables}
    current_page = 1
    paragraph_index = 0

    def _paragraph_has_page_break(para) -> bool:
        if para.paragraph_format and para.paragraph_format.page_break_before:
            return True
        return bool(para._element.xpath('.//*[local-name()="br" and @*[local-name()="type"]="page"]'))

    def _paragraph_has_section_break(para) -> bool:
        return bool(para._element.xpath('./*[local-name()="pPr"]/*[local-name()="sectPr"]'))

    def _append_paragraph_block(para, *, page_number: int) -> None:
        nonlocal paragraph_index
        text = para.text.strip()
        if not text:
            return
        paragraph_index += 1

        style_name = para.style.name if para.style else ""
        heading_level = _detect_heading_level(text)
        block_type = "heading" if heading_level else "caption" if _looks_like_caption(text) else "paragraph"

        run_font_sizes = []
        run_font_names = []
        for run in para.runs:
            if run.font and run.font.size:
                run_font_sizes.append(round(float(run.font.size.pt), 2))
            if run.font and run.font.name:
                run_font_names.append(run.font.name)
        font_size = run_font_sizes[0] if run_font_sizes else None
        font_name = Counter(run_font_names).most_common(1)[0][0] if run_font_names else ""
        if font_size:
            font_sizes.append(font_size)
        if font_name:
            font_names.append(font_name)

        line_spacing = None
        spacing_before = None
        spacing_after = None
        if para.paragraph_format:
            if para.paragraph_format.line_spacing:
                try:
                    line_spacing = float(para.paragraph_format.line_spacing)
                    line_spacings.append(line_spacing)
                except Exception:
                    line_spacing = None
            if para.paragraph_format.space_before:
                spacing_before = round(float(para.paragraph_format.space_before.pt), 2)
            if para.paragraph_format.space_after:
                spacing_after = round(float(para.paragraph_format.space_after.pt), 2)

        block = {
            "block_id": f"p-{paragraph_index}",
            "type": block_type,
            "page": page_number,
            "paragraph_index": paragraph_index,
            "heading_level": heading_level,
            "text": text,
            "bbox": None,
            "style": {
                "style_name": style_name,
                "font_name": font_name,
                "font_size_pt": font_size,
                "line_spacing": line_spacing,
                "space_before_pt": spacing_before,
                "space_after_pt": spacing_after,
                "alignment": para.alignment,
            },
        }
        blocks.append(block)
        anchors.append(
            {
                "anchor_id": f"anchor-p-{paragraph_index}",
                "block_id": block["block_id"],
                "page": page_number,
                "paragraph_index": paragraph_index,
                "bbox": None,
                "text_excerpt": text[:200],
            }
        )

        if block_type == "caption":
            captions.append(
                {
                    "caption_id": f"caption-{len(captions) + 1}",
                    "block_id": block["block_id"],
                    "label": text[:80],
                    "page": page_number,
                    "bbox": None,
                }
            )
            if re.match(r"^(图|Figure)", text, re.IGNORECASE):
                figures.append({"figure_id": f"fig-{len(figures) + 1}", "caption_id": captions[-1]["caption_id"], "page": page_number})
            if re.match(r"^(表|Table)", text, re.IGNORECASE):
                tables.append({"table_id": f"tbl-{len(tables) + 1}", "caption_id": captions[-1]["caption_id"], "page": page_number})

    def _append_table_block(table, *, page_number: int, table_index: int) -> None:
        nonlocal paragraph_index
        row_texts = []
        for row in table.rows:
            cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()) for cell in row.cells]
            text = " | ".join(item.strip() for item in cells if item and item.strip()).strip()
            if text:
                row_texts.append(text)
        table_text = "\n".join(row_texts).strip()
        if not table_text:
            return
        paragraph_index += 1
        block = {
            "block_id": f"tbl-{table_index}",
            "type": "table",
            "page": page_number,
            "paragraph_index": paragraph_index,
            "heading_level": None,
            "text": table_text,
            "bbox": None,
            "style": {"rows": len(table.rows), "columns": len(table.columns)},
        }
        blocks.append(block)
        anchors.append(
            {
                "anchor_id": f"anchor-tbl-{table_index}",
                "block_id": block["block_id"],
                "page": page_number,
                "paragraph_index": paragraph_index,
                "bbox": None,
                "text_excerpt": table_text[:200],
            }
        )
        tables.append(
            {
                "table_id": f"tbl-grid-{table_index}",
                "rows": len(table.rows),
                "columns": len(table.columns),
                "caption_id": None,
                "page": page_number,
            }
        )

    table_index = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p" and id(child) in para_map:
            para = para_map[id(child)]
            if para.paragraph_format and para.paragraph_format.page_break_before and blocks:
                current_page += 1
            _append_paragraph_block(para, page_number=current_page)
            if _paragraph_has_page_break(para) or _paragraph_has_section_break(para):
                current_page += 1
        elif tag == "tbl" and id(child) in table_map:
            table_index += 1
            _append_table_block(table_map[id(child)], page_number=current_page, table_index=table_index)

    page_total = max([block.get("page") or 1 for block in blocks] + [1])
    template_section = doc.sections[0] if doc.sections else None
    pages = []
    for page_index in range(1, page_total + 1):
        pages.append(
            {
                "page_number": page_index,
                "width_pt": round(float(template_section.page_width.pt), 2) if template_section and template_section.page_width else None,
                "height_pt": round(float(template_section.page_height.pt), 2) if template_section and template_section.page_height else None,
                "margins_cm": {
                    "top": round(float(template_section.top_margin.cm), 3) if template_section and template_section.top_margin else None,
                    "bottom": round(float(template_section.bottom_margin.cm), 3) if template_section and template_section.bottom_margin else None,
                    "left": round(float(template_section.left_margin.cm), 3) if template_section and template_section.left_margin else None,
                    "right": round(float(template_section.right_margin.cm), 3) if template_section and template_section.right_margin else None,
                },
            }
        )

    return {
        "text": "\n".join(block["text"] for block in blocks if block["text"]),
        "blocks": blocks,
        "anchors": anchors,
        "pages": pages,
        "figures": figures,
        "tables": tables,
        "captions": captions,
        "abstract": _extract_section_text_from_blocks(
            blocks,
            ["摘要", "摘 要", "abstract"],
            ["关键词", "key words", "keywords", "英文摘要", "外文摘要", "abstract", "目录", "引言", "绪论"],
        ),
        "format_metrics": {
            "font_name_summary": Counter(font_names).most_common(5),
            "font_size_summary_pt": Counter(font_sizes).most_common(5),
            "line_spacing_summary": Counter(line_spacings).most_common(5),
            "page_margin_summary_cm": pages[0]["margins_cm"] if pages else {},
        },
    }


def _parse_pdf_layout(path: Path) -> Dict[str, Any]:
    try:
        import fitz
    except Exception:
        text = extract_text_from_pdf(path)
        blocks = [
            {
                "block_id": f"p-{index}",
                "type": "heading" if _detect_heading_level(line) else "paragraph",
                "page": None,
                "paragraph_index": index,
                "heading_level": _detect_heading_level(line),
                "text": line.strip(),
                "bbox": None,
                "style": {},
            }
            for index, line in enumerate(text.splitlines(), start=1)
            if line.strip()
        ]
        anchors = [
            {
                "anchor_id": f"anchor-p-{index}",
                "block_id": item["block_id"],
                "page": None,
                "paragraph_index": index,
                "bbox": None,
                "text_excerpt": item["text"][:200],
            }
            for index, item in enumerate(blocks, start=1)
        ]
        return {
            "text": text,
            "blocks": blocks,
            "anchors": anchors,
            "pages": [],
            "figures": [],
            "tables": [],
            "captions": [],
            "abstract": _extract_section_text_from_blocks(
                blocks,
                ["摘要", "摘 要", "abstract"],
                ["关键词", "key words", "keywords", "英文摘要", "外文摘要", "目录", "引言", "绪论"],
            ),
            "format_metrics": {},
        }

    blocks: List[Dict[str, Any]] = []
    anchors: List[Dict[str, Any]] = []
    pages: List[Dict[str, Any]] = []
    figures: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []
    captions: List[Dict[str, Any]] = []
    text_chunks: List[str] = []

    with fitz.open(str(path)) as doc:
        for page_number, page in enumerate(doc, start=1):
            page_dict = page.get_text("dict")
            page_width = round(float(page.rect.width), 2)
            page_height = round(float(page.rect.height), 2)
            page_blocks = page_dict.get("blocks") or []
            text_blocks = []
            xs0: List[float] = []
            xs1: List[float] = []
            ys0: List[float] = []
            ys1: List[float] = []

            for block_index, item in enumerate(page_blocks, start=1):
                bbox = item.get("bbox")
                if item.get("type") == 1:
                    figures.append(
                        {
                            "figure_id": f"fig-{len(figures) + 1}",
                            "page": page_number,
                            "bbox": bbox,
                            "caption_id": None,
                        }
                    )
                    continue

                lines = item.get("lines") or []
                spans = []
                text_parts = []
                for line in lines:
                    for span in line.get("spans") or []:
                        span_text = str(span.get("text") or "").strip()
                        if span_text:
                            spans.append(span)
                            text_parts.append(span_text)
                text = " ".join(text_parts).strip()
                if not text:
                    continue

                font_size = round(float(max((span.get("size") or 0) for span in spans)), 2) if spans else None
                font_name = Counter(span.get("font") for span in spans if span.get("font")).most_common(1)
                font_name_value = font_name[0][0] if font_name else ""
                heading_level = _detect_heading_level(text)
                block_type = "heading" if heading_level else "caption" if _looks_like_caption(text) else "paragraph"
                block = {
                    "block_id": f"pdf-{page_number}-{block_index}",
                    "type": block_type,
                    "page": page_number,
                    "paragraph_index": len(blocks) + 1,
                    "heading_level": heading_level,
                    "text": text,
                    "bbox": bbox,
                    "style": {
                        "font_name": font_name_value,
                        "font_size_pt": font_size,
                    },
                }
                blocks.append(block)
                text_blocks.append(block)
                text_chunks.append(text)
                anchors.append(
                    {
                        "anchor_id": f"anchor-{block['block_id']}",
                        "block_id": block["block_id"],
                        "page": page_number,
                        "paragraph_index": block["paragraph_index"],
                        "bbox": bbox,
                        "text_excerpt": text[:200],
                    }
                )

                if bbox:
                    xs0.append(float(bbox[0]))
                    ys0.append(float(bbox[1]))
                    xs1.append(float(bbox[2]))
                    ys1.append(float(bbox[3]))

                if block_type == "caption":
                    captions.append(
                        {
                            "caption_id": f"caption-{len(captions) + 1}",
                            "block_id": block["block_id"],
                            "page": page_number,
                            "bbox": bbox,
                            "label": text[:80],
                        }
                    )
                    if re.match(r"^(图|Figure)", text, re.IGNORECASE):
                        figures.append({"figure_id": f"fig-{len(figures) + 1}", "page": page_number, "bbox": None, "caption_id": captions[-1]["caption_id"]})
                    if re.match(r"^(表|Table)", text, re.IGNORECASE):
                        tables.append({"table_id": f"tbl-{len(tables) + 1}", "page": page_number, "bbox": None, "caption_id": captions[-1]["caption_id"]})

            margins = None
            if xs0 and xs1 and ys0 and ys1:
                margins = {
                    "top": round(min(ys0), 2),
                    "bottom": round(page_height - max(ys1), 2),
                    "left": round(min(xs0), 2),
                    "right": round(page_width - max(xs1), 2),
                }

            pages.append(
                {
                    "page_number": page_number,
                    "width_pt": page_width,
                    "height_pt": page_height,
                    "margins_pt": margins,
                    "block_count": len(text_blocks),
                }
            )

    return {
        "text": "\n".join(text_chunks),
        "blocks": blocks,
        "anchors": anchors,
        "pages": pages,
        "figures": figures,
        "tables": tables,
        "captions": captions,
        "abstract": _extract_section_text_from_blocks(
            blocks,
            ["摘要", "摘 要", "abstract"],
            ["关键词", "key words", "keywords", "英文摘要", "外文摘要", "目录", "引言", "绪论"],
        ),
        "format_metrics": {
            "page_margin_summary_pt": pages[0]["margins_pt"] if pages else {},
        },
    }


def _parse_text_layout(path: Path, file_type: str) -> Dict[str, Any]:
    text = _read_text(path, file_type)
    blocks = []
    anchors = []
    figures = []
    tables = []
    captions = []

    for index, line in enumerate(text.splitlines(), start=1):
        stripped = line.strip()
        if not stripped:
            continue
        heading_level = _detect_heading_level(stripped)
        block_type = "heading" if heading_level else "caption" if _looks_like_caption(stripped) else "paragraph"
        block = {
            "block_id": f"text-{index}",
            "type": block_type,
            "page": None,
            "paragraph_index": index,
            "heading_level": heading_level,
            "text": stripped,
            "bbox": None,
            "style": {},
        }
        blocks.append(block)
        anchors.append(
            {
                "anchor_id": f"anchor-text-{index}",
                "block_id": block["block_id"],
                "page": None,
                "paragraph_index": index,
                "bbox": None,
                "text_excerpt": stripped[:200],
            }
        )
        if block_type == "caption":
            captions.append({"caption_id": f"caption-{len(captions) + 1}", "block_id": block["block_id"], "page": None, "bbox": None, "label": stripped[:80]})
            if re.match(r"^(图|Figure)", stripped, re.IGNORECASE):
                figures.append({"figure_id": f"fig-{len(figures) + 1}", "caption_id": captions[-1]["caption_id"]})
            if re.match(r"^(表|Table)", stripped, re.IGNORECASE):
                tables.append({"table_id": f"tbl-{len(tables) + 1}", "caption_id": captions[-1]["caption_id"]})

    return {
        "text": text,
        "blocks": blocks,
        "anchors": anchors,
        "pages": [],
        "figures": figures,
        "tables": tables,
        "captions": captions,
        "abstract": _extract_section_text_from_blocks(
            blocks,
            ["摘要", "摘 要", "abstract"],
            ["关键词", "key words", "keywords", "英文摘要", "外文摘要", "目录", "引言", "绪论"],
        ),
        "format_metrics": {},
    }


def build_evidence_bundle(
    paper_path: str | Path,
    *,
    template_name: Optional[str] = None,
    review_track: Optional[str] = None,
) -> Dict[str, Any]:
    path = Path(paper_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"论文文件不存在: {path}")

    file_type = detect_file_type(path)
    if file_type == "docx":
        parsed = _parse_docx_layout(path)
    elif file_type == "pdf":
        parsed = _parse_pdf_layout(path)
        grobid_payload = GrobidClient().process_fulltext_document(path)
        parsed = _merge_grobid_pdf_parse(parsed, grobid_payload)
    else:
        parsed = _parse_text_layout(path, file_type)

    text = parsed.get("text", "")
    if review_track in {"undergraduate", "graduate"}:
        track = review_track
        reasons = [f"前端/接口显式指定审查轨道: {review_track}"]
    else:
        track, reasons = _detect_review_track(text, template_name=template_name or "", paper_path=path)

    structure = parsed.get("grobid", {}).get("document_structure") or _build_structure_from_blocks(parsed.get("blocks", []), text)
    title = parsed.get("title") or _infer_title(path, text)

    engine = ReferenceEngine()
    refs = []
    stats = None
    references: List[Dict[str, Any]] = []
    citations: List[Dict[str, Any]] = []
    if file_type == "pdf" and parsed.get("grobid", {}).get("enabled"):
        references = parsed.get("grobid", {}).get("bibliography", []) or []
        citations = parsed.get("grobid", {}).get("citations", []) or []
        stats = _compute_reference_stats_from_records(references, citations)
    else:
        try:
            refs = engine.extract_from_paper(str(path))
        except Exception as exc:
            logger.warning("提取参考文献失败: %s", exc)
        try:
            stats = engine.validate(str(path), refs=refs)
        except Exception as exc:
            logger.warning("验证参考文献失败: %s", exc)
        citation_text = ReferenceParser.extract_body_text(str(path)) if file_type == "docx" else text
        citations = _extract_citation_contexts(citation_text)
        references = [_serialize_reference(ref, index) for index, ref in enumerate(refs, start=1)]

    layout_ast = {
        "ast_version": "layout_ast.v2",
        "nodes": parsed.get("blocks", []),
        "pages": parsed.get("pages", []),
        "anchors": parsed.get("anchors", []),
        "figures": parsed.get("figures", []),
        "tables": parsed.get("tables", []),
        "captions": parsed.get("captions", []),
        "document_structure": structure,
    }

    return {
        "bundle_version": "evidence_bundle.v2",
        "paper_id": path.stem,
        "source_path": str(path),
        "source_file_name": path.name,
        "file_type": file_type,
        "title": title,
        "language": _detect_language(text),
        "review_track": track,
        "routing_reasons": reasons,
        "text_length": len(text),
        "raw_text_excerpt": text[:18000],
        "document_structure": structure,
        "layout_ast": layout_ast,
        "layout": {
            "pages": parsed.get("pages", []),
            "blocks": parsed.get("blocks", []),
            "figures": parsed.get("figures", []),
            "tables": parsed.get("tables", []),
            "captions": parsed.get("captions", []),
        },
        "format_metrics": parsed.get("format_metrics", {}),
        "bibliography": {
            "references": references,
            "citations": citations,
            "reference_stats": stats if isinstance(stats, dict) else _serialize_reference_stats(stats),
        },
        "source_snippet_anchors": parsed.get("anchors", []),
        "parse_meta": {
            "grobid_enabled": bool(parsed.get("grobid", {}).get("enabled")),
            "grobid_source": parsed.get("grobid", {}).get("source"),
            "parser_profile": "docx_native" if file_type == "docx" else "pdf_layout_plus_grobid" if file_type == "pdf" else "text_fallback",
        },
    }


def build_section_digest(bundle: Dict[str, Any], *, max_sections: int = 16) -> List[Dict[str, Any]]:
    digest = []
    for node in (bundle.get("document_structure") or [])[:max_sections]:
        digest.append(
            {
                "section_title": node.get("heading", ""),
                "section_role": "",
                "key_points": [node.get("text_excerpt", "")[:180]] if node.get("text_excerpt") else [],
                "heading_level": node.get("level"),
                "page_start": node.get("page_start"),
                "page_end": node.get("page_end"),
            }
        )
    return digest


def build_evidence_index(bundle: Dict[str, Any], *, max_items: int = 64) -> List[Dict[str, Any]]:
    items = []
    for anchor in (bundle.get("source_snippet_anchors") or [])[:max_items]:
        items.append(
            {
                "evidence_id": anchor.get("anchor_id"),
                "evidence_type": "source_anchor",
                "locator": anchor.get("text_excerpt", ""),
                "page": anchor.get("page"),
                "block_id": anchor.get("block_id"),
            }
        )
    return items
