"""
格式检查工具 — 在 Harness 工具层运行的本地格式规则引擎

零 token 成本 — 完全本地执行。
"""
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from article_check.rules.latex.checker import LaTeXChecker
from article_check.rules.docx.checker import DocxChecker
from article_check.utils.file_utils import extract_text_from_docx, extract_text_from_pdf, read_paper_content

logger = logging.getLogger(__name__)
RULE_PROFILE_DIR = Path(__file__).resolve().parents[3] / "北师大论文格式要求"
UNDERGRAD_RULE_PATH = RULE_PROFILE_DIR / "bnu_undergraduate_template_rule_profile.json"
GRAD_RULE_PATH = RULE_PROFILE_DIR / "bnu_graduate_requirement_rule_profile.json"

# 全局检查器实例
_latex_checker = LaTeXChecker()
_docx_checker = DocxChecker()

SECTION_ALIASES = {
    "cover": ["封面", "论文封面"],
    "title_page": ["题名页"],
    "copyright": ["版权页"],
    "abstract": ["摘要", "摘 要", "abstract"],
    "en_abstract": ["英文摘要", "外文摘要", "abstract"],
    "keywords": ["关键词", "key words", "keywords"],
    "catalog": ["目录"],
    "introduction": ["引言", "前言", "绪论", "导论", "introduction"],
    "body": ["正文"],
    "related work": ["相关研究", "文献综述", "研究现状", "related work"],
    "method": ["研究方法", "方法", "方法设计", "系统框架", "建模", "method"],
    "experiment": ["实验", "实验设计", "实验设置", "实验评估", "experiment"],
    "result": ["结果", "实验结果", "结果分析", "results"],
    "discussion": ["讨论", "分析与讨论", "discussion"],
    "conclusion": ["结论", "总结", "结语", "conclusion"],
    "reference": ["参考文献", "references", "bibliography"],
}

TRACK_SECTION_REQUIREMENTS = {
    "undergraduate": [
        "cover", "abstract", "en_abstract", "keywords", "catalog",
        "introduction", "body", "conclusion", "reference",
    ],
    "graduate": [
        "cover", "title_page", "copyright", "abstract", "en_abstract",
        "keywords", "catalog", "introduction", "body", "reference",
    ],
}


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _load_text_for_structure(path: Path, file_type: str) -> str:
    if file_type == "docx":
        return extract_text_from_docx(path)
    if file_type == "pdf":
        return extract_text_from_pdf(path)
    return read_paper_content(path)


def _matches_section(section_name: str, normalized_text: str) -> bool:
    aliases = SECTION_ALIASES.get(section_name, [section_name])
    return any(_normalize_text(alias) in normalized_text for alias in aliases)


def _load_rule_profile(review_track: Optional[str]) -> Dict[str, Any]:
    path = UNDERGRAD_RULE_PATH if review_track == "undergraduate" else GRAD_RULE_PATH
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("读取规则画像失败 [%s]: %s", review_track, exc)
        return {}


def _resolve_expected_sections(expected_sections: Optional[List[str]], review_track: Optional[str]) -> List[str]:
    if expected_sections:
        return expected_sections
    if review_track in TRACK_SECTION_REQUIREMENTS:
        return TRACK_SECTION_REQUIREMENTS[review_track]
    return [
        "abstract", "introduction", "related work",
        "method", "experiment", "result",
        "discussion", "conclusion", "reference",
    ]


def _detect_cover_presence(path: Path, file_type: str) -> str:
    """返回 present / uncertain / missing。"""
    try:
        if file_type == "docx":
            from docx import Document

            doc = Document(str(path))
            lines: List[str] = []
            for para in doc.paragraphs[:24]:
                text = para.text.strip()
                if text:
                    lines.append(text)
                if para._element.xpath('.//*[local-name()="br" and @*[local-name()="type"]="page"]'):
                    break
            for table in doc.tables[:2]:
                for row in table.rows:
                    row_text = " ".join(" ".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()) for cell in row.cells).strip()
                    if row_text:
                        lines.append(row_text)
            joined = "\n".join(lines)
        elif file_type == "pdf":
            try:
                import fitz
                with fitz.open(str(path)) as doc:
                    joined = doc.load_page(0).get_text("text") if doc.page_count else ""
            except Exception:
                joined = extract_text_from_pdf(path).split("\f")[0]
        else:
            joined = _load_text_for_structure(path, file_type)
    except Exception:
        joined = _load_text_for_structure(path, file_type)

    normalized = _normalize_text(joined)
    if any(marker in normalized for marker in ["诚信承诺书", "论文使用授权说明"]):
        return "uncertain"

    title_hits = len(re.findall(r"[\u4e00-\u9fff]{4,}", joined))
    field_hits = sum(
        1 for marker in ["大学", "学院", "论文", "作者", "姓名", "学号", "指导教师", "导师", "专业"]
        if marker in joined
    )
    if field_hits >= 3 and title_hits >= 2:
        return "present"
    if field_hits >= 2:
        return "uncertain"
    return "missing"


def _detect_statement_presence(normalized_text: str) -> bool:
    return any(marker in normalized_text for marker in [
        "诚信承诺书", "论文使用授权说明", "原创性声明", "使用授权说明", "版权声明",
    ])


def check_latex_format(
    file_path: str,
    rules_filter: Optional[List[int]] = None,
) -> List[Dict[str, Any]]:
    """
    检查 LaTeX 文件格式

    零 token 成本 — 使用本地 chktex 规则引擎。

    Args:
        file_path: LaTeX 文件路径
        rules_filter: 可选，仅检查指定规则编号

    Returns:
        格式问题列表
    """
    logger.info(f"check_latex_format: {file_path}")
    issues = _latex_checker.check(file_path)

    if rules_filter:
        issues = [i for i in issues if i.get("rule_id") in rules_filter]

    return issues


def check_docx_format(
    file_path: str,
    template_path: Optional[str] = None,
    review_track: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    检查 Word 文档格式

    零 token 成本 — 使用 python-docx 规则引擎。

    Args:
        file_path: Word 文件路径
        template_path: 可选模板路径

    Returns:
        格式问题列表
    """
    logger.info(f"check_docx_format: {file_path}")
    checker = DocxChecker(
        template_path=template_path,
        review_track=review_track or "auto",
        rule_profile=_load_rule_profile(review_track or "graduate"),
    )
    issues = checker.check(file_path)
    return issues


def check_structure(
    file_path: str,
    file_type: str,
    expected_sections: Optional[List[str]] = None,
    review_track: Optional[str] = None,
) -> Dict[str, Any]:
    """
    检查论文结构完整性

    Args:
        file_path: 论文文件路径
        file_type: 文件类型 (latex/docx)
        expected_sections: 期望的章节列表

    Returns:
        结构检查结果
    """
    sections = _resolve_expected_sections(expected_sections, review_track)
    path = Path(file_path)
    detected_type = file_type or ("docx" if path.suffix.lower() == ".docx" else "latex")
    text = _load_text_for_structure(path, detected_type)
    normalized_text = _normalize_text(text)

    found = []
    missing = []
    cover_state = _detect_cover_presence(path, detected_type) if "cover" in sections else "missing"
    for sec in sections:
        if sec == "cover":
            if cover_state == "present":
                found.append(sec)
                continue
            if cover_state == "uncertain":
                found.append(f"{sec}:manual_check")
                continue
            missing.append(sec)
            continue
        if sec in {"copyright", "title_page"} and _detect_statement_presence(normalized_text):
            found.append(sec)
            continue
        if _matches_section(sec, normalized_text):
            found.append(sec)
        else:
            missing.append(sec)

    return {
        "issues": [
            *(
                [
                    {
                        "type": "cover_manual_confirmation",
                        "severity": "minor",
                        "section": "cover",
                        "description": "首页疑似为封面，但自动识别不够稳定，需人工确认。",
                        "suggestion": "请人工确认第一页是否为封面，并检查题名、作者、导师等字段是否完整。",
                    }
                ]
                if cover_state == "uncertain"
                else []
            ),
            *[
                {
                    "type": "missing_section",
                    "severity": "major" if s in ["abstract", "reference", "cover", "title_page", "copyright"] else "minor",
                    "section": s,
                    "description": f"缺少 '{s}' 章节",
                    "suggestion": f"请补充 {s} 相关内容并显式设置章节标题",
                }
                for s in missing
            ],
        ],
        "found_sections": found,
        "missing_sections": missing,
        "complete": len(missing) == 0,
        "review_track": review_track or "auto",
    }
